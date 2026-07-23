"""Renderer generique bloc -> style (V3). Un Document (liste de blocs, vocabulaire
commun) se rend dans n'importe quel template — un jeu de styles (couleurs, polices,
marges, disposition), jamais une structure de contenu — module par le palier
d'abonnement de l'utilisateur. Remplace les 8 anciens templates un-fichier-un-design
qui relisaient chacun les champs d'un schema Content fige.

AUCUNE regex, AUCUN parsing texte : uniquement les champs types des blocs Pydantic."""

import io
from dataclasses import dataclass, replace
from pathlib import Path
from typing import Literal

from docx.document import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor
from docx.table import _Cell

from app.core.document_engine.blocks import Block
from app.core.document_engine.document import Document as EngineDocument
from app.core.document_engine.templates._style_utils import (
    BLANC,
    GRIS_BANDE,
    GRIS_TEXTE,
    add_bottom_border,
    add_left_border,
    add_page_number_field,
    add_table_of_contents,
    new_document,
    remove_table_borders,
    set_cell_background,
    set_cell_margins,
    set_paragraph_shading,
    set_table_borders,
)

Container = DocxDocument | _Cell  # tout ce qui expose .add_paragraph() (Document python-docx ou cellule)
Richness = Literal["sobre", "soigne", "premium"]


class RendererError(Exception):
    pass


def _tier_richness(tier: str) -> Richness:
    """introduction/goutte -> sobre, source/fleuve -> soigne, ocean -> premium. Ces
    quatre outils ne sont de toute facon jamais disponibles en introduction (voir
    app/core/llm/router.py) ; le cas est traite ici par prudence, pas par necessite."""
    if tier in ("introduction", "goutte"):
        return "sobre"
    if tier in ("source", "fleuve"):
        return "soigne"
    return "premium"


@dataclass
class TemplateStyle:
    doc_type: str
    label: str
    cv_sidebar: bool = False  # CV : colonne laterale sombre (moderne) vs une seule colonne (classique)
    letter_banner: bool = False  # Lettre : bandeau colore en en-tete (moderne) vs sobre (standard)
    # Chaque template pro_doc/academic a desormais un style de page de garde qui lui est
    # propre et systematique (jamais absente) — cv/cover_letter n'ont pas ce champ car ils
    # n'ont pas vocation a en avoir une (voir DOCUMENT_SCHEMAS, blocks.py).
    cover_page_style: Literal["banner", "minimal", "formal", "clean"] | None = None
    numbered_headings: bool = False
    reliure_margin: bool = False
    margins_cm: tuple[float, float, float, float] = (2.5, 2.5, 2.5, 2.5)  # gauche, droite, haut, bas
    font_name: str = "Arial"  # jamais Calibri — Arial (moderne) ou Times New Roman (classique/formel)
    # Couleurs propres au template — defaut = bleu/marine Boulga (templates pro_doc/
    # academic, inchanges). cv/cover_letter diversifient desormais ces teintes par
    # contexte (etudiant, academique, concours...) plutot que de partager toujours
    # les 2 memes couleurs de marque.
    accent_hex: str = "1565C0"
    dark_hex: str = "0B1F3A"


TEMPLATE_STYLES: dict[str, TemplateStyle] = {
    "cv_modern": TemplateStyle(doc_type="cv", label="Professionnel", cv_sidebar=True, font_name="Arial"),
    "cv_classic": TemplateStyle(
        doc_type="cv",
        label="Étudiant / Scolaire",
        cv_sidebar=False,
        font_name="Arial",
        accent_hex="0E7C6B",
        dark_hex="0B4A3D",
    ),
    "cv_academique": TemplateStyle(
        doc_type="cv",
        label="Académique / Universitaire",
        cv_sidebar=False,
        font_name="Times New Roman",
        accent_hex="333333",
        dark_hex="1A1A1A",
    ),
    "cv_concours": TemplateStyle(
        doc_type="cv",
        label="Concours / Administratif",
        cv_sidebar=False,
        font_name="Times New Roman",
        accent_hex="37474F",
        dark_hex="263238",
    ),
    "letter_standard": TemplateStyle(
        doc_type="cover_letter", label="Standard", letter_banner=False, font_name="Times New Roman"
    ),
    "letter_modern": TemplateStyle(
        doc_type="cover_letter", label="Moderne", letter_banner=True, font_name="Arial"
    ),
    "letter_concours": TemplateStyle(
        doc_type="cover_letter",
        label="Concours / Fonction publique",
        letter_banner=False,
        font_name="Times New Roman",
        accent_hex="37474F",
        dark_hex="263238",
    ),
    "letter_academique": TemplateStyle(
        doc_type="cover_letter",
        label="Académique / Recherche",
        letter_banner=False,
        font_name="Times New Roman",
        accent_hex="333333",
        dark_hex="1A1A1A",
    ),
    "pro_corporate": TemplateStyle(
        doc_type="pro_doc",
        label="Corporate",
        cover_page_style="banner",
        numbered_headings=True,
        font_name="Arial",
    ),
    "pro_minimal": TemplateStyle(
        doc_type="pro_doc",
        label="Minimal",
        cover_page_style="minimal",
        font_name="Arial",
        accent_hex="45607A",
        dark_hex="37474F",
    ),
    "pro_moderne": TemplateStyle(
        doc_type="pro_doc",
        label="Moderne",
        cover_page_style="banner",
        numbered_headings=False,
        font_name="Arial",
        accent_hex="0E7C6B",
        dark_hex="0B4A3D",
    ),
    "academic_formal": TemplateStyle(
        doc_type="academic",
        label="Formel",
        cover_page_style="formal",
        reliure_margin=True,
        margins_cm=(3, 2, 2.5, 2.5),
        font_name="Times New Roman",
    ),
    "academic_clean": TemplateStyle(
        doc_type="academic",
        label="Épuré",
        cover_page_style="clean",
        font_name="Arial",
        accent_hex="37474F",
        dark_hex="263238",
    ),
    "academic_classique": TemplateStyle(
        doc_type="academic",
        label="Classique",
        cover_page_style="clean",
        font_name="Times New Roman",
        accent_hex="7B2D26",
        dark_hex="4A1B17",
    ),
}


def available_templates(doc_type: str | None = None) -> list[str]:
    if doc_type is None:
        return sorted(TEMPLATE_STYLES.keys())
    return sorted(name for name, style in TEMPLATE_STYLES.items() if style.doc_type == doc_type)


def _add_picture_safe(container: Container, photo_bytes: bytes, width_cm: float, alignment) -> None:
    """Insere une photo/logo dans son propre paragraphe — jamais bloquant : une image
    corrompue ou dans un format que python-docx ne sait pas lire ne doit jamais faire
    echouer tout le rendu du document (voir documents.py render_document, qui
    telecharge ces octets depuis le bucket Storage avant d'appeler render)."""
    try:
        p = container.add_paragraph()
        p.alignment = alignment
        p.add_run().add_picture(io.BytesIO(photo_bytes), width=Cm(width_cm))
    except Exception:
        pass


# --- Socle commun : rendu dans le document directement (jamais dans une cellule, pour
# beneficier des styles Heading natifs de Word, necessaires au champ TOC) -----------


def _render_heading(doc: DocxDocument, block, style: TemplateStyle, richness: Richness, number: str | None) -> None:
    text = f"{number}. {block.text}" if number else block.text
    heading = doc.add_heading(text, level=min(block.level, 4))
    if not heading.runs:
        return
    run = heading.runs[0]
    run.font.color.rgb = (
        RGBColor.from_string(style.dark_hex)
        if richness == "sobre" or block.level > 1
        else RGBColor.from_string(style.accent_hex)
    )
    if richness != "sobre" and block.level == 1:
        add_bottom_border(heading, style.accent_hex if richness == "premium" else style.dark_hex)


def _render_table(container: Container, block, richness: Richness, style: TemplateStyle) -> None:
    table = container.add_table(rows=1, cols=max(len(block.headers), 1))
    set_table_borders(table)
    header_color = style.dark_hex if richness == "sobre" else style.accent_hex
    for cell, header in zip(table.rows[0].cells, block.headers):
        set_cell_background(cell, header_color)
        set_cell_margins(cell, top_cm=0.2, left_cm=0.25, bottom_cm=0.2, right_cm=0.25)
        cell.text = header
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = BLANC
    for row_idx, row in enumerate(block.rows):
        cells = table.add_row().cells
        for cell, value in zip(cells, row):
            set_cell_margins(cell, top_cm=0.2, left_cm=0.25, bottom_cm=0.2, right_cm=0.25)
            if row_idx % 2 == 1:
                set_cell_background(cell, GRIS_BANDE)
            cell.text = value
    if block.caption:
        caption_p = container.add_paragraph()
        caption_run = caption_p.add_run(block.caption)
        caption_run.italic = True
        caption_run.font.size = Pt(9)


def _render_quote(container: Container, block, style: TemplateStyle) -> None:
    p = container.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    set_paragraph_shading(p, "F2F5FA")
    add_left_border(p, style.accent_hex, size=18)
    run = p.add_run(block.text)
    run.italic = True


def _render_toc(doc: DocxDocument, richness: Richness, style: TemplateStyle) -> None:
    toc_heading = doc.add_paragraph()
    toc_run = toc_heading.add_run("SOMMAIRE" if richness != "sobre" else "Sommaire")
    toc_run.font.bold = True
    if richness != "sobre":
        toc_run.font.color.rgb = RGBColor.from_string(style.accent_hex)
    add_table_of_contents(doc)
    doc.add_page_break()


def _render_bibliography(doc: DocxDocument, block, style: TemplateStyle) -> None:
    heading = doc.add_heading("Bibliographie", level=1)
    if heading.runs:
        heading.runs[0].font.color.rgb = RGBColor.from_string(style.dark_hex)
    for entry in block.entries:
        doc.add_paragraph(entry, style="List Bullet")


def _render_cover_page_banner(
    doc: DocxDocument, block, richness: Richness, style: TemplateStyle, photo_bytes: bytes | None
) -> None:
    """pro_corporate/pro_moderne : bandeau colore pleine largeur, poursuit directement
    dans le contenu (pas de saut de page — le bandeau tient lieu d'en-tete de premiere
    page)."""
    header_table = doc.add_table(rows=1, cols=1)
    remove_table_borders(header_table)
    cell = header_table.rows[0].cells[0]
    set_cell_background(cell, style.accent_hex if richness != "sobre" else style.dark_hex)
    set_cell_margins(cell, top_cm=0.5, left_cm=0.8, bottom_cm=0.5, right_cm=0.8)
    cell.paragraphs[0].text = ""
    title_p = cell.paragraphs[0]
    title_run = title_p.add_run(block.title)
    title_run.font.size = Pt(20)
    title_run.font.bold = True
    title_run.font.color.rgb = BLANC

    meta_bits = [block.author, block.institution, block.date, *block.extra.values()]
    meta_p = cell.add_paragraph()
    meta_run = meta_p.add_run(" | ".join(b for b in meta_bits if b))
    meta_run.font.size = Pt(10)
    meta_run.font.color.rgb = BLANC
    meta_p.paragraph_format.space_after = Pt(14)
    doc.add_paragraph()
    if photo_bytes:
        _add_picture_safe(doc, photo_bytes, 3.5, WD_ALIGN_PARAGRAPH.CENTER)


def _render_cover_page_minimal(doc: DocxDocument, block, style: TemplateStyle, photo_bytes: bytes | None) -> None:
    """pro_minimal : page de titre dediee mais sobre — pas de bloc couleur, juste le
    titre et les metadonnees centres, sur leur propre page (au lieu de n'avoir aucune
    page de garde comme avant : chaque template pro_doc en a desormais une)."""
    for _ in range(6):
        doc.add_paragraph()
    if photo_bytes:
        _add_picture_safe(doc, photo_bytes, 3, WD_ALIGN_PARAGRAPH.CENTER)
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run(block.title)
    title_run.font.size = Pt(22)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor.from_string(style.dark_hex)

    meta_bits = [block.author, block.institution, block.date, *block.extra.values()]
    meta_line = " · ".join(b for b in meta_bits if b)
    if meta_line:
        meta_p = doc.add_paragraph()
        meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta_p.paragraph_format.space_before = Pt(10)
        meta_run = meta_p.add_run(meta_line)
        meta_run.font.size = Pt(11)
        meta_run.font.color.rgb = GRIS_TEXTE
    doc.add_page_break()


def _render_cover_page_academic(doc: DocxDocument, block, style: TemplateStyle, photo_bytes: bytes | None) -> None:
    """academic_formal / academic_clean / academic_classique : page de garde
    academique classique centree — seule differe la marge de reliure, deja geree au
    niveau section (style.margins_cm), pas ici."""
    for _ in range(4):
        doc.add_paragraph()
    if block.institution:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(block.institution).font.size = Pt(13)
        doc.add_paragraph()
    if photo_bytes:
        _add_picture_safe(doc, photo_bytes, 3, WD_ALIGN_PARAGRAPH.CENTER)
        doc.add_paragraph()

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run(block.title)
    title_run.font.size = Pt(20)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor.from_string(style.dark_hex)

    for _ in range(4):
        doc.add_paragraph()
    if block.author:
        author_p = doc.add_paragraph()
        author_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        author_p.add_run(f"Présenté par {block.author}").font.size = Pt(12)
    if block.supervisor:
        supervisor_p = doc.add_paragraph()
        supervisor_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        supervisor_p.add_run(f"Sous la direction de {block.supervisor}").font.size = Pt(11)

    for _ in range(3):
        doc.add_paragraph()
    if block.date:
        date_p = doc.add_paragraph()
        date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_p.add_run(block.date).font.size = Pt(11)

    doc.add_page_break()


def _render_cover_page(
    doc: DocxDocument, block, style: TemplateStyle, richness: Richness, photo_bytes: bytes | None
) -> None:
    if style.cover_page_style == "banner":
        _render_cover_page_banner(doc, block, richness, style, photo_bytes)
    elif style.cover_page_style == "minimal":
        _render_cover_page_minimal(doc, block, style, photo_bytes)
    else:  # "formal" ou "clean" — meme composition academique
        _render_cover_page_academic(doc, block, style, photo_bytes)


# --- Blocs lettre --------------------------------------------------------------------


def _render_letter_header(doc: DocxDocument, block, style: TemplateStyle) -> None:
    if style.letter_banner:
        header_table = doc.add_table(rows=1, cols=1)
        remove_table_borders(header_table)
        cell = header_table.rows[0].cells[0]
        set_cell_background(cell, style.dark_hex)
        set_cell_margins(cell, top_cm=0.4, left_cm=0.8, bottom_cm=0.4, right_cm=0.8)
        cell.paragraphs[0].text = ""
        name_run = cell.paragraphs[0].add_run(block.sender_name)
        name_run.font.size = Pt(18)
        name_run.font.bold = True
        name_run.font.color.rgb = BLANC
        contact_p = cell.add_paragraph()
        contact_run = contact_p.add_run(" | ".join(block.sender_contact))
        contact_run.font.size = Pt(9.5)
        contact_run.font.color.rgb = BLANC
        cell.paragraphs[0].paragraph_format.space_before = Pt(10)
        contact_p.paragraph_format.space_after = Pt(10)
        doc.add_paragraph()
    else:
        name_p = doc.add_paragraph()
        name_p.add_run(block.sender_name).bold = True
        for line in block.sender_contact:
            doc.add_paragraph(line)
        doc.add_paragraph()

    recipient_lines = [x for x in [block.recipient_name, block.recipient_title, block.company_name] if x]
    for line in recipient_lines:
        doc.add_paragraph(line)

    if block.place or block.date:
        date_p = doc.add_paragraph()
        date_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        date_p.add_run(", ".join(x for x in [block.place, block.date] if x))


_LETTER_TIER_TAGLINES: dict[Richness, str | None] = {
    "sobre": None,
    "soigne": "boulga.ai",
    "premium": "Généré avec Boulga AI",
}


def _render_subject(doc: DocxDocument, block, style: TemplateStyle, richness: Richness) -> None:
    doc.add_paragraph()
    p = doc.add_paragraph()
    # L'accent colore s'applique des que le template le demande (bandeau moderne) OU
    # que le palier n'est pas le plus sobre — les deux axes (template, palier) restent
    # independants mais se cumulent visuellement.
    if style.letter_banner or richness != "sobre":
        run = p.add_run(block.text)
        run.font.bold = True
        run.font.color.rgb = RGBColor.from_string(style.accent_hex)
        run.font.size = Pt(12)
    else:
        run = p.add_run(f"Objet : {block.text}")
        run.bold = True
    doc.add_paragraph()
    tagline = _LETTER_TIER_TAGLINES[richness]
    if tagline:
        tag_p = doc.add_paragraph()
        tag_run = tag_p.add_run(tagline)
        tag_run.font.size = Pt(7)
        tag_run.font.color.rgb = RGBColor.from_string(style.accent_hex)
        doc.add_paragraph()


def _render_signature(doc: DocxDocument, block) -> None:
    doc.add_paragraph()
    doc.add_paragraph(block.closing)
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run(block.name)


# --- Blocs CV : fonctionnent aussi bien dans le document (classique) que dans une
# cellule de tableau (moderne, colonne principale ou laterale) ----------------------


def _white(run) -> None:
    run.font.color.rgb = BLANC


def _bullet_indent(paragraph) -> None:
    """Retrait suspendu (puce a gauche, texte aligne en dessous meme sur plusieurs
    lignes) — les puces manuelles ("• texte" en un seul run) n'ont sinon aucun retrait
    de continuation et une ligne qui wrap revient sous la puce, pas sous le texte."""
    paragraph.paragraph_format.left_indent = Cm(0.45)
    paragraph.paragraph_format.first_line_indent = Cm(-0.45)


def _render_contact(
    container: Container,
    block,
    richness: Richness,
    style: TemplateStyle,
    *,
    white_text: bool,
    photo_bytes: bytes | None = None,
) -> None:
    dark = RGBColor.from_string(style.dark_hex)
    accent = RGBColor.from_string(style.accent_hex)

    if photo_bytes:
        # Centree dans la colonne laterale (sidebar, etroite) ; alignee a droite dans
        # une mise en page classique une colonne, pour laisser le nom/titre a gauche.
        _add_picture_safe(
            container, photo_bytes, 2.8, WD_ALIGN_PARAGRAPH.CENTER if white_text else WD_ALIGN_PARAGRAPH.RIGHT
        )

    name_p = container.add_paragraph()
    name_run = name_p.add_run(block.full_name or "—")
    name_run.font.size = Pt(18 if white_text else 20)
    name_run.font.bold = True
    if white_text:
        _white(name_run)
    else:
        name_run.font.color.rgb = dark

    if block.title:
        title_p = container.add_paragraph()
        title_run = title_p.add_run(block.title)
        title_run.font.size = Pt(12)
        if white_text:
            _white(title_run)
        else:
            # Accent colore reserve aux paliers au-dessus du plus sobre — cf. modulation
            # par palier (V3-5) : le CV classique reste volontairement plus neutre en
            # Goutte, plus marque en Source/Fleuve/Ocean.
            title_run.font.color.rgb = dark if richness == "sobre" else accent

    bits = [block.email, block.phone, block.address, block.linkedin]
    for bit in (b for b in bits if b):
        p = container.add_paragraph()
        run = p.add_run(bit)
        run.font.size = Pt(9.5)
        if white_text:
            _white(run)

    # Etat civil (cv_concours notamment) : ignore silencieusement si absent, jamais
    # requis pour les autres templates.
    civil_bits = []
    if block.birth_date or block.birth_place:
        parts = [p for p in [block.birth_date and f"Né(e) le {block.birth_date}", block.birth_place] if p]
        civil_bits.append(" à ".join(parts) if len(parts) == 2 else parts[0])
    if block.nationality:
        civil_bits.append(f"Nationalité : {block.nationality}")
    for line in civil_bits:
        p = container.add_paragraph()
        run = p.add_run(line)
        run.font.size = Pt(9.5)
        if white_text:
            _white(run)


def _render_summary(container: Container, block, *, white_text: bool) -> None:
    p = container.add_paragraph()
    run = p.add_run(block.text)
    run.italic = True
    run.font.size = Pt(10)
    if white_text:
        _white(run)


def _render_experience(container: Container, block, *, white_text: bool) -> None:
    period = " – ".join(x for x in [block.start, block.end or ("présent" if block.start else "")] if x)
    p = container.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run(block.position + (f" — {block.company}" if block.company else ""))
    run.font.bold = True
    run.font.size = Pt(10.5)
    if white_text:
        _white(run)
    if period or block.location:
        meta_p = container.add_paragraph()
        meta_run = meta_p.add_run(period + (f" | {block.location}" if block.location else ""))
        meta_run.italic = True
        meta_run.font.size = Pt(9)
        if white_text:
            _white(meta_run)
    if block.description:
        desc_p = container.add_paragraph()
        desc_run = desc_p.add_run(block.description)
        desc_run.font.size = Pt(10)
        if white_text:
            _white(desc_run)
    for achievement in block.achievements:
        bullet_p = container.add_paragraph()
        _bullet_indent(bullet_p)
        bullet_run = bullet_p.add_run(f"• {achievement}")
        bullet_run.font.size = Pt(9.5)
        if white_text:
            _white(bullet_run)


def _render_education(container: Container, block, *, white_text: bool) -> None:
    p = container.add_paragraph()
    run = p.add_run(block.degree + (f" — {block.institution}" if block.institution else ""))
    run.font.bold = True
    run.font.size = Pt(10.5)
    if white_text:
        _white(run)
    if block.year or block.location:
        meta_p = container.add_paragraph()
        meta_run = meta_p.add_run(block.year + (f" | {block.location}" if block.location else ""))
        meta_run.italic = True
        meta_run.font.size = Pt(9)
        if white_text:
            _white(meta_run)
    if block.details:
        details_p = container.add_paragraph()
        details_run = details_p.add_run(block.details)
        details_run.font.size = Pt(9.5)
        if white_text:
            _white(details_run)


def _render_skill_group(container: Container, block, *, white_text: bool) -> None:
    label_p = container.add_paragraph()
    label_p.paragraph_format.space_before = Pt(10)
    label_run = label_p.add_run((block.label or "Compétences").upper())
    label_run.font.size = Pt(10)
    label_run.font.bold = True
    if white_text:
        _white(label_run)
    for skill in block.skills:
        p = container.add_paragraph()
        _bullet_indent(p)
        run = p.add_run(f"• {skill}")
        run.font.size = Pt(9.5)
        if white_text:
            _white(run)


def _render_language_group(container: Container, block, *, white_text: bool) -> None:
    label_p = container.add_paragraph()
    label_p.paragraph_format.space_before = Pt(10)
    label_run = label_p.add_run("LANGUES")
    label_run.font.size = Pt(10)
    label_run.font.bold = True
    if white_text:
        _white(label_run)
    for lang in block.languages:
        p = container.add_paragraph()
        run = p.add_run(f"{lang.language} — {lang.level}" if lang.level else lang.language)
        run.font.size = Pt(9.5)
        if white_text:
            _white(run)


def _render_heading_in_container(container: Container, block, style: TemplateStyle, *, white_text: bool) -> None:
    """Variante sans style Heading natif — utilisee dans les cellules (CV moderne),
    ou .add_heading() n'existe pas. Pas de TOC concernee pour un CV."""
    p = container.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    run = p.add_run(block.text.upper())
    run.font.bold = True
    run.font.size = Pt(max(14 - block.level * 2, 10))
    if white_text:
        _white(run)
    else:
        run.font.color.rgb = RGBColor.from_string(style.accent_hex)


def _render_paragraph_in_container(container: Container, block, *, white_text: bool) -> None:
    p = container.add_paragraph()
    run = p.add_run(block.text)
    run.font.size = Pt(10)
    if white_text:
        _white(run)


def _render_bullet_list_in_container(container: Container, block, *, white_text: bool) -> None:
    for item in block.items:
        p = container.add_paragraph()
        _bullet_indent(p)
        run = p.add_run(f"• {item}")
        run.font.size = Pt(9.5)
        if white_text:
            _white(run)


def _render_bibliography_in_container(container: Container, block, *, white_text: bool) -> None:
    """Section 'Publications' du CV academique — reutilise le bloc bibliography deja
    existant (pro_doc/academic) plutot que d'inventer un type de bloc dedie."""
    label_p = container.add_paragraph()
    label_p.paragraph_format.space_before = Pt(10)
    label_run = label_p.add_run("PUBLICATIONS")
    label_run.font.size = Pt(10)
    label_run.font.bold = True
    if white_text:
        _white(label_run)
    for entry in block.entries:
        p = container.add_paragraph()
        _bullet_indent(p)
        run = p.add_run(f"• {entry}")
        run.font.size = Pt(9.5)
        if white_text:
            _white(run)


# Renderers qui ne dependent pas du palier — "contact" est special-case a part car son
# accent colore varie avec la richesse (voir _dispatch_cv_block).
CV_BLOCK_RENDERERS = {
    "summary": _render_summary,
    "experience": _render_experience,
    "education": _render_education,
    "skill_group": _render_skill_group,
    "language_group": _render_language_group,
    "paragraph": _render_paragraph_in_container,
    "bullet_list": _render_bullet_list_in_container,
    "bibliography": _render_bibliography_in_container,
}

# Modulation par palier (V3-5) pour le CV : un signe discret et honnete de la richesse
# du palier plutot qu'une refonte complete de la mise en page — goutte reste sobre,
# source/fleuve ajoutent une petite touche, ocean une mention complete.
_CV_TIER_TAGLINES: dict[Richness, str | None] = {
    "sobre": None,
    "soigne": "boulga.ai",
    "premium": "Généré avec Boulga AI — palier Océan",
}


def _append_cv_tagline(container: Container, richness: Richness, style: TemplateStyle, *, white_text: bool) -> None:
    text = _CV_TIER_TAGLINES[richness]
    if not text:
        return
    p = container.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    run = p.add_run(text)
    run.font.size = Pt(7)
    if white_text:
        _white(run)
    else:
        run.font.color.rgb = RGBColor.from_string(style.accent_hex)


def _dispatch_cv_block(
    container: Container,
    block,
    richness: Richness,
    style: TemplateStyle,
    *,
    white_text: bool,
    photo_bytes: bytes | None = None,
) -> None:
    if block.type == "contact":
        _render_contact(container, block, richness, style, white_text=white_text, photo_bytes=photo_bytes)
        return
    if block.type == "heading":
        _render_heading_in_container(container, block, style, white_text=white_text)
        return
    renderer = CV_BLOCK_RENDERERS.get(block.type)
    if renderer is not None:
        renderer(container, block, white_text=white_text)


def _render_cv_classic(
    doc: DocxDocument, blocks: list[Block], richness: Richness, style: TemplateStyle, photo_bytes: bytes | None
) -> None:
    for block in blocks:
        _dispatch_cv_block(doc, block, richness, style, white_text=False, photo_bytes=photo_bytes)
    _append_cv_tagline(doc, richness, style, white_text=False)


def _render_cv_sidebar(
    doc: DocxDocument, blocks: list[Block], richness: Richness, style: TemplateStyle, photo_bytes: bytes | None
) -> None:
    section = doc.sections[0]
    section.left_margin = Cm(0)
    section.right_margin = Cm(0)
    section.top_margin = Cm(0)
    section.bottom_margin = Cm(0)

    table = doc.add_table(rows=1, cols=2)
    remove_table_borders(table)
    left, right = table.rows[0].cells
    table.columns[0].width = Cm(6)
    table.columns[1].width = Cm(15)
    left.width = Cm(6)
    right.width = Cm(15)
    set_cell_background(left, style.dark_hex)
    # Marges internes de cellule : la section entiere est a marge zero (le tableau doit
    # occuper toute la page), donc sans ceci le texte touche directement les bords de
    # page/cellule — seul rempart contre ce collage, vu qu'aucune marge de section
    # n'existe plus a ce stade.
    set_cell_margins(left, top_cm=0.9, left_cm=0.7, bottom_cm=0.9, right_cm=0.7)
    set_cell_margins(right, top_cm=0.9, left_cm=1.0, bottom_cm=0.9, right_cm=0.8)
    for cell in (left, right):
        for p in cell.paragraphs:
            p.paragraph_format.space_after = Pt(4)
    left.paragraphs[0].text = ""
    right.paragraphs[0].text = ""

    sidebar_types = {"contact", "skill_group", "language_group"}
    for block in blocks:
        target, white_text = (left, True) if block.type in sidebar_types else (right, False)
        _dispatch_cv_block(target, block, richness, style, white_text=white_text, photo_bytes=photo_bytes)
    _append_cv_tagline(left, richness, style, white_text=True)


# --- Dispatch generique (documents pro/academique/lettre — rendu direct dans doc) ----


def _render_linear(
    doc: DocxDocument, blocks: list[Block], style: TemplateStyle, richness: Richness, photo_bytes: bytes | None
) -> None:
    counters = [0, 0, 0, 0]
    for block in blocks:
        if block.type == "heading":
            number = None
            if style.numbered_headings:
                idx = min(block.level, 4) - 1
                counters[idx] += 1
                for deeper in range(idx + 1, 4):
                    counters[deeper] = 0
                number = ".".join(str(c) for c in counters[: idx + 1] if c)
            _render_heading(doc, block, style, richness, number)
        elif block.type == "paragraph":
            doc.add_paragraph(block.text)
        elif block.type == "bullet_list":
            for item in block.items:
                doc.add_paragraph(item, style="List Bullet")
        elif block.type == "numbered_list":
            for item in block.items:
                doc.add_paragraph(item, style="List Number")
        elif block.type == "table":
            _render_table(doc, block, richness, style)
        elif block.type == "quote":
            _render_quote(doc, block, style)
        elif block.type == "spacer":
            doc.add_paragraph()
        elif block.type == "page_break":
            doc.add_page_break()
        elif block.type == "cover_page":
            _render_cover_page(doc, block, style, richness, photo_bytes)
        elif block.type == "table_of_contents":
            _render_toc(doc, richness, style)
        elif block.type == "bibliography":
            _render_bibliography(doc, block, style)
        elif block.type == "letter_header":
            _render_letter_header(doc, block, style)
        elif block.type == "subject":
            _render_subject(doc, block, style, richness)
        elif block.type == "signature":
            _render_signature(doc, block)
        # les blocs CV (contact/summary/experience/...) n'apparaissent jamais dans un
        # document non-CV (filtres en amont par validate_document sur le vocabulaire).


def render(
    document: EngineDocument,
    template_name: str,
    tier: str,
    output_dir: Path,
    accent_override: str | None = None,
    dark_override: str | None = None,
    # Octets deja telecharges par l'appelant (voir documents.py render_document) a
    # partir du photo_path du bloc contact/cover_page concerne — le renderer ne fait
    # lui-meme aucun acces reseau/Storage, uniquement de l'embarquement d'image.
    photo_bytes: bytes | None = None,
) -> Path:
    style = TEMPLATE_STYLES.get(template_name)
    if style is None:
        raise RendererError(f"Template inconnu : {template_name}")
    if style.doc_type != document.doc_type:
        raise RendererError(
            f"Le template « {template_name} » ne s'applique pas au type « {document.doc_type} »."
        )
    # Couleurs choisies par le user (palette curatee, voir palette.py) : accent_override
    # couvre les titres/accents, dark_override le nom/les elements secondaires et,
    # pour les templates a sidebar/bandeau, leur fond colore — le fond blanc de la
    # page elle-meme n'est jamais concerne, uniquement ces deux teintes de style.
    if accent_override or dark_override:
        style = replace(
            style,
            accent_hex=accent_override or style.accent_hex,
            dark_hex=dark_override or style.dark_hex,
        )

    richness = _tier_richness(tier)
    doc = new_document(style.font_name)
    section = doc.sections[0]
    left, right, top, bottom = style.margins_cm
    section.left_margin = Cm(left)
    section.right_margin = Cm(right)
    section.top_margin = Cm(top)
    section.bottom_margin = Cm(bottom)

    if document.doc_type == "cv" and style.cv_sidebar:
        _render_cv_sidebar(doc, document.blocks, richness, style, photo_bytes)
    elif document.doc_type == "cv":
        _render_cv_classic(doc, document.blocks, richness, style, photo_bytes)
    else:
        _render_linear(doc, document.blocks, style, richness, photo_bytes)

    has_pagination_context = style.cover_page_style is not None or any(
        b.type == "table_of_contents" for b in document.blocks
    )
    if has_pagination_context:
        footer_p = section.footer.paragraphs[0]
        footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if richness == "premium":
            footer_p.add_run("Boulga AI — ").font.size = Pt(8)
        add_page_number_field(footer_p)

    output_path = output_dir / "output.docx"
    doc.save(str(output_path))
    return output_path
