"""CV classique : une colonne sobre, nom centre, sections separees par des lignes
horizontales. Pas de couleur de fond, seul le Bleu Boulga accentue les titres."""

from pathlib import Path

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from app.core.document_engine.renderer import register_template
from app.core.document_engine.schema import CVContent
from app.core.document_engine.templates._style_utils import (
    BLEU_BOULGA,
    MARINE,
    add_bottom_border,
    new_document,
)


def build(content: CVContent, output_dir: Path) -> Path:
    doc = new_document()
    section = doc.sections[0]
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)

    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_p.add_run(content.full_name)
    name_run.font.size = Pt(22)
    name_run.font.bold = True
    name_run.font.color.rgb = MARINE

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.add_run(content.title).font.size = Pt(13)

    contact_parts = [content.contact.email]
    if content.contact.phone:
        contact_parts.append(content.contact.phone)
    if content.contact.address:
        contact_parts.append(content.contact.address)
    if content.contact.linkedin:
        contact_parts.append(content.contact.linkedin)
    contact_p = doc.add_paragraph()
    contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_run = contact_p.add_run(" | ".join(contact_parts))
    contact_run.font.size = Pt(9.5)

    def section_heading(text: str) -> None:
        doc.add_paragraph()
        p = doc.add_paragraph()
        run = p.add_run(text.upper())
        run.font.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = BLEU_BOULGA
        add_bottom_border(p)

    if content.summary:
        section_heading("Profil")
        doc.add_paragraph(content.summary)

    if content.experiences:
        section_heading("Experience professionnelle")
        for exp in content.experiences:
            p = doc.add_paragraph()
            p.add_run(f"{exp.title} - {exp.company}").bold = True
            dates = exp.start_date + (f" - {exp.end_date}" if exp.end_date else " - present")
            meta = dates + (f" | {exp.location}" if exp.location else "")
            meta_p = doc.add_paragraph()
            meta_run = meta_p.add_run(meta)
            meta_run.italic = True
            meta_run.font.size = Pt(9.5)
            doc.add_paragraph(exp.description)
            for achievement in exp.achievements:
                bullet = doc.add_paragraph(style="List Bullet")
                bullet.add_run(achievement)

    if content.education:
        section_heading("Formation")
        for edu in content.education:
            p = doc.add_paragraph()
            p.add_run(f"{edu.degree} - {edu.institution}").bold = True
            meta = edu.year + (f" | {edu.location}" if edu.location else "")
            meta_p = doc.add_paragraph()
            meta_run = meta_p.add_run(meta)
            meta_run.italic = True
            meta_run.font.size = Pt(9.5)
            if edu.details:
                doc.add_paragraph(edu.details)

    if content.skills:
        section_heading("Competences")
        doc.add_paragraph(", ".join(content.skills))

    if content.languages:
        section_heading("Langues")
        doc.add_paragraph(", ".join(f"{l.language} ({l.level})" for l in content.languages))

    if content.certifications:
        section_heading("Certifications")
        for cert in content.certifications:
            doc.add_paragraph(cert, style="List Bullet")

    output_path = output_dir / "output.docx"
    doc.save(str(output_path))
    return output_path


register_template("cv_classic", build)
