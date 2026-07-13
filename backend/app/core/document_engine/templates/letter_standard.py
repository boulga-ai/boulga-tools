"""Lettre de motivation, format classique francais : expediteur, destinataire, lieu/date,
objet, corps, formule de politesse. Marges 2.5 cm."""

from pathlib import Path

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from app.core.document_engine.renderer import register_template
from app.core.document_engine.schema import CoverLetterContent
from app.core.document_engine.templates._style_utils import new_document


def build(content: CoverLetterContent, output_dir: Path) -> Path:
    doc = new_document()
    section = doc.sections[0]
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    sender_p = doc.add_paragraph()
    sender_p.add_run(content.full_name).bold = True
    contact_lines = [content.contact.email]
    if content.contact.phone:
        contact_lines.append(content.contact.phone)
    if content.contact.address:
        contact_lines.append(content.contact.address)
    for line in contact_lines:
        doc.add_paragraph(line)

    doc.add_paragraph()
    if content.recipient_name:
        doc.add_paragraph(content.recipient_name)
    if content.recipient_title:
        doc.add_paragraph(content.recipient_title)
    doc.add_paragraph(content.company_name)

    doc.add_paragraph()
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_p.add_run(content.date)

    doc.add_paragraph()
    subject_p = doc.add_paragraph()
    subject_run = subject_p.add_run(f"Objet : {content.subject}")
    subject_run.bold = True

    doc.add_paragraph()
    doc.add_paragraph(content.greeting)

    for paragraph in content.paragraphs:
        p = doc.add_paragraph(paragraph)
        p.paragraph_format.space_after = Pt(8)

    doc.add_paragraph()
    doc.add_paragraph(content.closing)
    doc.add_paragraph()
    signature_p = doc.add_paragraph()
    signature_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    signature_p.add_run(content.signature)

    output_path = output_dir / "output.docx"
    doc.save(str(output_path))
    return output_path


register_template("letter_standard", build)
