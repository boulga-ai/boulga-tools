"""Lettre de motivation moderne : bande marine en en-tete, mise en page aeree, accent
Bleu Boulga sur l'objet."""

from pathlib import Path

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from app.core.document_engine.renderer import register_template
from app.core.document_engine.schema import CoverLetterContent
from app.core.document_engine.templates._style_utils import (
    BLANC,
    BLEU_BOULGA,
    new_document,
    set_cell_background,
    remove_table_borders,
)


def build(content: CoverLetterContent, output_dir: Path) -> Path:
    doc = new_document()
    section = doc.sections[0]
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    section.top_margin = Cm(0)
    section.bottom_margin = Cm(2)

    header_table = doc.add_table(rows=1, cols=1)
    remove_table_borders(header_table)
    header_cell = header_table.rows[0].cells[0]
    set_cell_background(header_cell, "0B1F3A")
    header_cell.paragraphs[0].text = ""
    name_p = header_cell.paragraphs[0]
    name_run = name_p.add_run(content.full_name)
    name_run.font.size = Pt(18)
    name_run.font.bold = True
    name_run.font.color.rgb = BLANC
    contact_p = header_cell.add_paragraph()
    contact_bits = [content.contact.email]
    if content.contact.phone:
        contact_bits.append(content.contact.phone)
    contact_run = contact_p.add_run(" | ".join(contact_bits))
    contact_run.font.size = Pt(9.5)
    contact_run.font.color.rgb = BLANC
    header_cell.paragraphs[0].paragraph_format.space_before = Pt(10)
    contact_p.paragraph_format.space_after = Pt(10)

    doc.add_paragraph()

    recipient_lines = [x for x in [content.recipient_name, content.recipient_title, content.company_name] if x]
    for line in recipient_lines:
        doc.add_paragraph(line)

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_p.add_run(content.date)

    doc.add_paragraph()
    subject_p = doc.add_paragraph()
    subject_run = subject_p.add_run(content.subject)
    subject_run.font.bold = True
    subject_run.font.color.rgb = BLEU_BOULGA
    subject_run.font.size = Pt(12)

    doc.add_paragraph()
    doc.add_paragraph(content.greeting)

    for paragraph in content.paragraphs:
        p = doc.add_paragraph(paragraph)
        p.paragraph_format.space_after = Pt(10)

    doc.add_paragraph()
    doc.add_paragraph(content.closing)
    doc.add_paragraph()
    signature_p = doc.add_paragraph()
    signature_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    signature_run = signature_p.add_run(content.signature)
    signature_run.font.bold = True

    output_path = output_dir / "output.docx"
    doc.save(str(output_path))
    return output_path


register_template("letter_modern", build)
