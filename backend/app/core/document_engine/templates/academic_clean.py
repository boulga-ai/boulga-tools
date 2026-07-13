"""Document academique epure : page de garde simple, minimaliste, sans sommaire ni
couleurs marquees."""

from pathlib import Path

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from app.core.document_engine.renderer import register_template
from app.core.document_engine.schema import AcademicDocContent, OutlineSection
from app.core.document_engine.templates._style_utils import add_page_number_field, new_document

DOC_TYPE_LABELS = {"rapport_stage": "Rapport de stage", "memoire": "Memoire", "these": "These"}


def _add_outline_section(doc, section: OutlineSection, sections: dict[str, str]) -> None:
    doc.add_heading(section.title, level=min(section.level, 3))
    content = sections.get(section.id)
    if content:
        for block in content.split("\n\n"):
            if block.strip():
                doc.add_paragraph(block.strip())
    for child in section.children:
        _add_outline_section(doc, child, sections)


def build(content: AcademicDocContent, output_dir: Path) -> Path:
    doc = new_document()
    section = doc.sections[0]
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    for _ in range(6):
        doc.add_paragraph()
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.add_run(content.title).font.size = Pt(18)
    title_p.runs[0].font.bold = True

    type_p = doc.add_paragraph()
    type_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    type_p.add_run(DOC_TYPE_LABELS.get(content.doc_type, content.doc_type))

    doc.add_paragraph()
    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_bits = [content.author, content.institution, content.year]
    if content.supervisor:
        meta_bits.append(f"Dir. {content.supervisor}")
    meta_p.add_run(" | ".join(meta_bits)).font.size = Pt(10)

    doc.add_page_break()

    if content.abstract:
        doc.add_heading("Resume", level=1)
        doc.add_paragraph(content.abstract)
        doc.add_page_break()

    for outline_section in content.outline.sections:
        _add_outline_section(doc, outline_section, content.sections)

    if content.bibliography:
        doc.add_page_break()
        doc.add_heading("Bibliographie", level=1)
        for entry in content.bibliography:
            doc.add_paragraph(entry, style="List Bullet")

    footer_p = section.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number_field(footer_p)

    output_path = output_dir / "output.docx"
    doc.save(str(output_path))
    return output_path


register_template("academic_clean", build)
