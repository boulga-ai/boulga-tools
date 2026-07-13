"""Document professionnel minimal : sobre, sans couleurs, universel."""

from pathlib import Path

from docx.shared import Cm, Pt

from app.core.document_engine.renderer import register_template
from app.core.document_engine.schema import ProDocContent, ProSection
from app.core.document_engine.templates._style_utils import new_document


def _add_section(doc, section: ProSection, level: int) -> None:
    doc.add_heading(section.title, level=min(level, 3))
    if section.content:
        doc.add_paragraph(section.content)
    for sub in section.subsections:
        _add_section(doc, sub, level + 1)


def build(content: ProDocContent, output_dir: Path) -> Path:
    doc = new_document()
    section = doc.sections[0]
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    title_p = doc.add_paragraph()
    title_run = title_p.add_run(content.title)
    title_run.font.size = Pt(18)
    title_run.font.bold = True

    meta_bits = [content.organization, content.author, content.date]
    meta_p = doc.add_paragraph()
    meta_run = meta_p.add_run(" | ".join(b for b in meta_bits if b))
    meta_run.font.size = Pt(9.5)
    meta_run.italic = True

    doc.add_paragraph()

    for sec in content.sections:
        _add_section(doc, sec, level=1)

    output_path = output_dir / "output.docx"
    doc.save(str(output_path))
    return output_path


register_template("pro_minimal", build)
