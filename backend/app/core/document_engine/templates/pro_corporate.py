"""Document professionnel corporate : bande Bleu Boulga en en-tete avec organisation,
sommaire (TOC Word natif) si plus de 3 sections, pied de page 'Confidentiel'."""

from pathlib import Path

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from app.core.document_engine.renderer import register_template
from app.core.document_engine.schema import ProDocContent, ProSection
from app.core.document_engine.templates._style_utils import (
    BLANC,
    BLEU_BOULGA,
    MARINE,
    add_page_number_field,
    add_table_of_contents,
    new_document,
    remove_table_borders,
    set_cell_background,
)


def _add_section(doc, section: ProSection, level: int) -> None:
    heading = doc.add_heading(section.title, level=min(level, 3))
    run = heading.runs[0]
    run.font.color.rgb = BLEU_BOULGA if level == 1 else MARINE
    run.font.size = Pt(max(16 - level * 2, 11))

    if section.content:
        doc.add_paragraph(section.content)

    for sub in section.subsections:
        _add_section(doc, sub, level + 1)


def build(content: ProDocContent, output_dir: Path) -> Path:
    doc = new_document()
    section = doc.sections[0]
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    header_table = doc.add_table(rows=1, cols=1)
    remove_table_borders(header_table)
    header_cell = header_table.rows[0].cells[0]
    set_cell_background(header_cell, "1565C0")
    header_cell.paragraphs[0].text = ""
    title_p = header_cell.paragraphs[0]
    title_run = title_p.add_run(content.title)
    title_run.font.size = Pt(20)
    title_run.font.bold = True
    title_run.font.color.rgb = BLANC
    title_p.paragraph_format.space_before = Pt(14)

    meta_p = header_cell.add_paragraph()
    meta_bits = [content.organization, content.author, content.date]
    meta_run = meta_p.add_run(" | ".join(b for b in meta_bits if b))
    meta_run.font.size = Pt(10)
    meta_run.font.color.rgb = BLANC
    meta_p.paragraph_format.space_after = Pt(14)

    doc.add_paragraph()

    if len(content.sections) > 3:
        toc_heading = doc.add_paragraph()
        toc_run = toc_heading.add_run("SOMMAIRE")
        toc_run.font.bold = True
        toc_run.font.color.rgb = BLEU_BOULGA
        add_table_of_contents(doc)
        doc.add_page_break()

    for sec in content.sections:
        _add_section(doc, sec, level=1)

    footer = section.footer
    footer_p = footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_p.add_run("Confidentiel — ")
    footer_run.font.size = Pt(8)
    footer_run.font.italic = True
    add_page_number_field(footer_p)

    output_path = output_dir / "output.docx"
    doc.save(str(output_path))
    return output_path


register_template("pro_corporate", build)
