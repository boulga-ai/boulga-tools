"""Document academique formel : page de garde, table des matieres (champ Word natif),
numerotation de pages, marges 3 cm a gauche (reliure)."""

from pathlib import Path

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from app.core.document_engine.renderer import register_template
from app.core.document_engine.schema import AcademicDocContent, OutlineSection
from app.core.document_engine.templates._style_utils import (
    BLEU_BOULGA,
    MARINE,
    add_page_number_field,
    add_table_of_contents,
    new_document,
)

DOC_TYPE_LABELS = {"rapport_stage": "Rapport de stage", "memoire": "Memoire", "these": "These"}


def _add_outline_section(doc, section: OutlineSection, sections: dict[str, str]) -> None:
    heading = doc.add_heading(section.title, level=min(section.level, 3))
    heading.runs[0].font.color.rgb = BLEU_BOULGA if section.level == 1 else MARINE

    content = sections.get(section.id)
    if content:
        for block in content.split("\n\n"):
            if block.strip():
                doc.add_paragraph(block.strip())

    for child in section.children:
        _add_outline_section(doc, child, sections)


def build(content: AcademicDocContent, output_dir: Path) -> Path:
    doc = new_document()
    section = doc.sections[0]
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    # Page de garde
    for _ in range(4):
        doc.add_paragraph()
    institution_p = doc.add_paragraph()
    institution_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    institution_p.add_run(content.institution).font.size = Pt(13)

    doc.add_paragraph()
    type_p = doc.add_paragraph()
    type_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    type_run = type_p.add_run(DOC_TYPE_LABELS.get(content.doc_type, content.doc_type).upper())
    type_run.font.size = Pt(14)
    type_run.font.color.rgb = BLEU_BOULGA
    type_run.font.bold = True

    for _ in range(2):
        doc.add_paragraph()
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run(content.title)
    title_run.font.size = Pt(20)
    title_run.font.bold = True
    title_run.font.color.rgb = MARINE

    for _ in range(4):
        doc.add_paragraph()
    author_p = doc.add_paragraph()
    author_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author_p.add_run(f"Presente par {content.author}").font.size = Pt(12)

    if content.supervisor:
        supervisor_p = doc.add_paragraph()
        supervisor_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        supervisor_p.add_run(f"Sous la direction de {content.supervisor}").font.size = Pt(11)

    for _ in range(3):
        doc.add_paragraph()
    year_p = doc.add_paragraph()
    year_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    year_p.add_run(content.year).font.size = Pt(11)

    doc.add_page_break()

    if content.abstract:
        doc.add_heading("Resume", level=1)
        doc.add_paragraph(content.abstract)
        doc.add_page_break()

    toc_heading = doc.add_paragraph()
    toc_run = toc_heading.add_run("TABLE DES MATIERES")
    toc_run.font.bold = True
    toc_run.font.color.rgb = BLEU_BOULGA
    add_table_of_contents(doc)
    doc.add_page_break()

    for outline_section in content.outline.sections:
        _add_outline_section(doc, outline_section, content.sections)

    if content.bibliography:
        doc.add_page_break()
        doc.add_heading("Bibliographie", level=1)
        for entry in content.bibliography:
            doc.add_paragraph(entry, style="List Bullet")

    footer_p = section.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number_field(footer_p)

    output_path = output_dir / "output.docx"
    doc.save(str(output_path))
    return output_path


register_template("academic_formal", build)
