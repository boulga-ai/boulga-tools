"""Utilitaires partages par les templates python-docx : couleurs de la charte Boulga,
polices, et helpers OOXML pour ce que python-docx n'expose pas nativement (ombrage de
cellule, bordures de paragraphe)."""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

MARINE = RGBColor(0x0B, 0x1F, 0x3A)
BLEU_BOULGA = RGBColor(0x15, 0x65, 0xC0)
GRIS_TEXTE = RGBColor(0x33, 0x33, 0x33)
BLANC = RGBColor(0xFF, 0xFF, 0xFF)
GRIS_BANDE = "F2F2F2"  # fond des lignes alternees de tableau (tres clair, discret)


def set_cell_background(cell, hex_color: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_margins(cell, *, top_cm: float, left_cm: float, bottom_cm: float, right_cm: float) -> None:
    """Marges internes d'une cellule (w:tcMar), absentes de l'API haut-niveau de
    python-docx. Sans elles, le texte touche directement les bordures/bords de page
    (cas du CV moderne, dont la section entiere est a marge zero pour que la sidebar
    occupe toute la page — voir _render_cv_sidebar)."""
    tc_mar = OxmlElement("w:tcMar")
    for edge, value in (("top", top_cm), ("left", left_cm), ("bottom", bottom_cm), ("right", right_cm)):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:w"), str(Cm(value).twips))
        el.set(qn("w:type"), "dxa")
        tc_mar.append(el)
    cell._tc.get_or_add_tcPr().append(tc_mar)


def set_paragraph_shading(paragraph, hex_color: str) -> None:
    """Fond teinte derriere un paragraphe entier (w:shd sur pPr) — utilise pour les
    citations, faute d'API python-docx dediee (seul le fond de cellule est expose)."""
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), hex_color)
    paragraph._p.get_or_add_pPr().append(shading)


def add_left_border(paragraph, color_hex: str = "1565C0", size: int = 18) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(size))
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), color_hex)
    borders.append(left)
    p_pr.append(borders)


def set_table_borders(table, color_hex: str = "D9D9D9", size: int = 4) -> None:
    """Bordures fines et discretes (gris clair) plutot que le style Word par defaut,
    plus lourd visuellement — utilise pour les tableaux de contenu (blocs `table`)."""
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(size))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color_hex)
        borders.append(el)
    tbl_pr.append(borders)


def remove_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "nil")
        borders.append(el)
    tbl_pr.append(borders)


def add_bottom_border(paragraph, color_hex: str = "1565C0", size: int = 6) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color_hex)
    borders.append(bottom)
    p_pr.append(borders)


def new_document(font_name: str) -> Document:
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = font_name
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = GRIS_TEXTE
    # Le gabarit par defaut de python-docx donne au style Normal un espacement apres
    # paragraphe d'environ 8pt (herite du docx par defaut de la bibliotheque) — sans
    # l'annuler ici, CHAQUE paragraphe (chaque puce, chaque ligne de meta-info...)
    # traine cet espace en plus, ce qui donne un rendu beaucoup trop aere entre des
    # lignes censees rester proches (ex: puces d'une meme experience). Les espacements
    # volontaires entre sections restent : ils sont fixes explicitement par bloc
    # (space_before = Pt(6)/Pt(10) dans renderer.py), independamment de ce defaut.
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(2)
    # Les styles Heading du gabarit par defaut de python-docx referencent la police de
    # theme ("+Headings") plutot que le style Normal — sans ceci, les titres resteraient
    # dans la police du theme (Calibri Light) meme apres avoir change Normal.
    for i in range(1, 5):
        doc.styles[f"Heading {i}"].font.name = font_name
    return doc


def add_title(doc, text: str, size: int = 20, color: RGBColor = MARINE, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = True
    run.font.color.rgb = color
    return p


def add_section_heading(doc, text: str, color: RGBColor = BLEU_BOULGA):
    p = doc.add_paragraph()
    run = p.add_run(text.upper())
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = color
    add_bottom_border(p, "1565C0" if color == BLEU_BOULGA else "0B1F3A")
    return p


def add_table_of_contents(doc) -> None:
    """Insere un champ TOC Word natif (mis a jour par Word a l'ouverture ou via F9,
    comme dans le cahier des charges technique). python-docx n'expose pas cette API :
    construction manuelle des runs de champ (begin/instrText/separate/end)."""
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = 'TOC \\o "1-3" \\h \\z \\u'

    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")

    placeholder = OxmlElement("w:t")
    placeholder.text = "Mettre a jour le sommaire (F9)"

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_begin)
    run._r.append(instr_text)
    run._r.append(fld_separate)
    run._r.append(placeholder)
    run._r.append(fld_end)


def add_page_number_field(paragraph) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr_text)
    run._r.append(fld_end)
