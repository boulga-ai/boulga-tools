"""CV moderne : 2 colonnes, colonne gauche (30%) fond marine (contact/competences/
langues en blanc), colonne droite (70%) nom/titre/resume/experiences/formation.
Accent Bleu Boulga sur les titres de section."""

from pathlib import Path

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

from app.core.document_engine.renderer import register_template
from app.core.document_engine.schema import CVContent
from app.core.document_engine.templates._style_utils import (
    BLANC,
    BLEU_BOULGA,
    MARINE,
    new_document,
    remove_table_borders,
    set_cell_background,
)


def _white_heading(cell, text: str) -> None:
    p = cell.add_paragraph()
    run = p.add_run(text.upper())
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.color.rgb = BLANC
    p.space_after = Pt(2)


def _white_line(cell, text: str, size: int = 9.5) -> None:
    p = cell.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = BLANC


def build(content: CVContent, output_dir: Path) -> Path:
    doc = new_document()
    section = doc.sections[0]
    section.left_margin = Cm(0)
    section.right_margin = Cm(0)
    section.top_margin = Cm(0)
    section.bottom_margin = Cm(0)

    table = doc.add_table(rows=1, cols=2)
    remove_table_borders(table)
    left, right = table.rows[0].cells
    table.columns[0].width = Cm(6)
    table.columns[1].width = Cm(15)
    left.width = Cm(6)
    right.width = Cm(15)

    set_cell_background(left, "0B1F3A")
    for cell in (left, right):
        for p in cell.paragraphs:
            p.paragraph_format.space_after = Pt(4)

    left.paragraphs[0].text = ""
    _white_heading(left, "Contact")
    _white_line(left, content.contact.email)
    if content.contact.phone:
        _white_line(left, content.contact.phone)
    if content.contact.address:
        _white_line(left, content.contact.address)
    if content.contact.linkedin:
        _white_line(left, content.contact.linkedin)

    if content.skills:
        left.add_paragraph()
        _white_heading(left, "Competences")
        for skill in content.skills:
            _white_line(left, f"- {skill}")

    if content.languages:
        left.add_paragraph()
        _white_heading(left, "Langues")
        for lang in content.languages:
            _white_line(left, f"{lang.language} - {lang.level}")

    if content.certifications:
        left.add_paragraph()
        _white_heading(left, "Certifications")
        for cert in content.certifications:
            _white_line(left, f"- {cert}")

    right.paragraphs[0].text = ""
    name_p = right.paragraphs[0]
    name_run = name_p.add_run(content.full_name)
    name_run.font.size = Pt(20)
    name_run.font.bold = True
    name_run.font.color.rgb = MARINE

    title_p = right.add_paragraph()
    title_run = title_p.add_run(content.title)
    title_run.font.size = Pt(13)
    title_run.font.color.rgb = BLEU_BOULGA

    if content.summary:
        right.add_paragraph()
        summary_p = right.add_paragraph()
        summary_run = summary_p.add_run(content.summary)
        summary_run.font.size = Pt(10)
        summary_run.italic = True

    def right_heading(text: str) -> None:
        right.add_paragraph()
        p = right.add_paragraph()
        run = p.add_run(text.upper())
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = BLEU_BOULGA

    if content.experiences:
        right_heading("Experience professionnelle")
        for exp in content.experiences:
            p = right.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            run = p.add_run(f"{exp.title} - {exp.company}")
            run.font.bold = True
            run.font.size = Pt(10.5)
            dates = exp.start_date + (f" - {exp.end_date}" if exp.end_date else " - present")
            meta_p = right.add_paragraph()
            meta_run = meta_p.add_run(dates + (f" | {exp.location}" if exp.location else ""))
            meta_run.font.size = Pt(9)
            meta_run.italic = True
            desc_p = right.add_paragraph()
            desc_p.add_run(exp.description).font.size = Pt(10)
            for achievement in exp.achievements:
                ach_p = right.add_paragraph()
                ach_p.paragraph_format.left_indent = Cm(0.5)
                ach_p.add_run(f"- {achievement}").font.size = Pt(9.5)

    if content.education:
        right_heading("Formation")
        for edu in content.education:
            p = right.add_paragraph()
            run = p.add_run(f"{edu.degree} - {edu.institution}")
            run.font.bold = True
            run.font.size = Pt(10.5)
            meta_p = right.add_paragraph()
            meta_run = meta_p.add_run(edu.year + (f" | {edu.location}" if edu.location else ""))
            meta_run.font.size = Pt(9)
            meta_run.italic = True
            if edu.details:
                right.add_paragraph().add_run(edu.details).font.size = Pt(9.5)

    output_path = output_dir / "output.docx"
    doc.save(str(output_path))
    return output_path


register_template("cv_modern", build)
